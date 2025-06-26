import os
import time
import signal
import uuid
import tempfile
import asyncio
from datetime import datetime, timedelta
from typing import Dict, Optional, List, Any
from concurrent.futures import ThreadPoolExecutor
import hashlib
import subprocess
# Fix for OpenMP runtime conflict - must be at the very top
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request, BackgroundTasks
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import mimetypes

from sqlalchemy.ext.asyncio import AsyncSession
from contextlib import asynccontextmanager

from database.database import engine, AsyncSessionLocal
from database.models import Base, User
from database.crud import create_user, authenticate_user

import httpx
import json

# RAG imports
import faiss
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import PyPDF2
from docx import Document as DocxDocument
import io

# Global variables
current_model = 'llama3.1:8b'
embedding_model = 'nomic-embed-text'  # Default Ollama embedding model
ollama_process = None
thread_pool = ThreadPoolExecutor(max_workers=4)
MODEL_PATH='http://localhost:8080'
EMBEDDING_PATH='http://localhost:8081'

# In-memory session storage
class RAGSession:
    def __init__(self, session_id: str):
        self.session_id = session_id
        self.created_at = datetime.now()
        self.last_accessed = datetime.now()
        self.document_name: Optional[str] = None
        self.chunks: List[Document] = []
        self.embeddings: Optional[np.ndarray] = None
        self.vector_store: Optional[faiss.IndexFlatIP] = None
        self.chunk_texts: List[str] = []
        
    def update_access_time(self):
        self.last_accessed = datetime.now()
        
    def has_document(self) -> bool:
        return self.document_name is not None
        
    def clear_document_data(self):
        """Clear all document-related data"""
        self.document_name = None
        self.chunks = []
        self.embeddings = None
        self.vector_store = None
        self.chunk_texts = []

# Global session storage
rag_sessions: Dict[str, RAGSession] = {}
SESSION_TIMEOUT_HOURS = 2

async def cleanup_expired_sessions():
    """Remove expired sessions"""
    current_time = datetime.now()
    expired_sessions = []
    
    for session_id, session in rag_sessions.items():
        if current_time - session.last_accessed > timedelta(hours=SESSION_TIMEOUT_HOURS):
            expired_sessions.append(session_id)
    
    for session_id in expired_sessions:
        del rag_sessions[session_id]
        print(f"Expired session removed: {session_id}")

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Create database tables at startup
    try:
        async with engine.begin() as conn:
            await conn.run_sync(Base.metadata.create_all)
    except Exception as e:
        print(f"Database initialization error: {e}")
    
    # Start cleanup task
    cleanup_task = asyncio.create_task(periodic_cleanup())
    
    yield
    
    # Cleanup on shutdown
    cleanup_task.cancel()
    rag_sessions.clear()

async def periodic_cleanup():
    """Periodically clean up expired sessions"""
    while True:
        try:
            await asyncio.sleep(3600)  # Run every hour
            await cleanup_expired_sessions()
        except asyncio.CancelledError:
            break
        except Exception as e:
            print(f"Error in periodic cleanup: {e}")

app = FastAPI(lifespan=lifespan)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    return {"message": "FastAPI Login Server with Stateless RAG is running"}

class LoginRequest(BaseModel):
    username: str
    password: str

class LoginResponse(BaseModel):
    message: str
    username: str

@app.post("/login", response_model=LoginResponse)
async def login(request: LoginRequest):
    """Authenticate user with username and password"""
    username = request.username.strip()
    password = request.password

    async with AsyncSessionLocal() as db:
        user = await authenticate_user(db, username, password)
        
        if not user:
            raise HTTPException(
                status_code=401,
                detail="Invalid username or password"
            )
        
        return LoginResponse(
            message="Login successful",
            username=username
        )

class SwitchModelRequest(BaseModel):
    modelName: str

@app.post("/switch-model")
async def switch_model(request: SwitchModelRequest):
    global current_model

    model_id = request.modelName.strip()
    
    if not model_id:
        return JSONResponse({
            "status": "error",
            "message": "Model name cannot be empty"
        }, status_code=400)

    try:
        # Check if Ollama is running
        async with httpx.AsyncClient(timeout=5.0) as client:
            try:
                await client.get("http://localhost:11434/api/tags")
            except Exception:
                return JSONResponse({
                    "status": "error",
                    "message": "Ollama server is not running"
                }, status_code=503)
                
        # Update current model
        current_model = model_id

        return JSONResponse({
            "status": "success",
            "message": f"Successfully switched to model: {model_id}",
            "current_model": current_model
        })

    except Exception as e:
        print(f"Error switching model: {e}")
        return JSONResponse({
            "status": "error",
            "message": f"Failed to switch model: {str(e)}"
        }, status_code=500)

class MessageRequest(BaseModel):
    chatHistory: str
    message: str
    model : str
    ragStatus: bool
    session_id: str

class MessageResponse(BaseModel):
    message: str

@app.post("/message", response_model=MessageResponse)
async def send_message(request: MessageRequest):
    """Send a message to the current model"""
    history = request.chatHistory
    prompt = request.message
    if not prompt.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")
    
    final_prompt = f"""You are a helpful AI assistant 'V'. Provide accurate, concise, and engaging responses.
GUIDELINES:
Be conversational and friendly while staying professional
Give direct answers
CHAT HISTORY:
{history if history else "No previous conversation"}

CURRENT USER MESSAGE:
{prompt}

Respond helpfully to the user's message, referencing previous conversation when relevant."""
    
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            try:
                await client.get("http://localhost:11434/api/tags", timeout=5.0)
            except Exception:
                raise HTTPException(status_code=503, detail="Ollama server is not responding")
            
            response = await client.post(
                url="http://localhost:11434/api/generate",
                json={
                    "model": current_model,
                    "prompt": final_prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.3,
                        "top_p": 0.95,
                        "max_tokens": 1000,
                        "stop": ["<SUF>", "<PRE>", "</PRE>", "</SUF>", "< EOT >", "\\end", "<MID>", "</MID>", "##"]
                    }
                },
                timeout=60.0
            )
            
            if response.status_code != 200:
                raise HTTPException(
                    status_code=500, 
                    detail=f"Ollama API returned status {response.status_code}"
                )
            
            data = response.json()
            
            if "response" not in data:
                raise HTTPException(
                    status_code=500, 
                    detail="Invalid response format from Ollama"
                )
            
            ai_response = data["response"].strip()
            
            if not ai_response:
                ai_response = "I apologize, but I couldn't generate a response. Please try again."
            
            return MessageResponse(message=ai_response)

    except HTTPException:
        raise
    except Exception as e:
        print(f"Ollama Error: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail=f"Error generating response: {str(e)}"
        )

async def gen_context(session_id, query):
    """Query the RAG system for a specific session"""
    if session_id not in rag_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = rag_sessions[session_id]
    session.update_access_time()
    
    if not session.has_document():
        raise HTTPException(
            status_code=400,
            detail="No document found in session. Please upload a document first."
        )

    if not query:
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    
    try:
        # Get query embedding
        print(f"Getting embedding for query: {query}")
        query_embedding = await get_embeddings([query])
        print(f"Query embedding shape: {query_embedding.shape}")
        
        # Ensure embedding is float32 and normalized
        query_embedding = query_embedding.astype(np.float32)
        faiss.normalize_L2(query_embedding)
        
        # Validate vector store
        if session.vector_store is None:
            raise HTTPException(status_code=500, detail="Vector store not initialized")
        
        # Search similar chunks
        k = min(4, len(session.chunks))
        print(f"Searching for {k} similar chunks")
        
        # Perform the search
        scores, indices = session.vector_store.search(query_embedding, k)
        print(f"Search completed. Scores: {scores}, Indices: {indices}")
        
        # Get relevant chunks
        relevant_chunks = []
        sources = set()
        
        for i, idx in enumerate(indices[0]):
            if 0 <= idx < len(session.chunks):  # Validate index bounds
                chunk = session.chunks[idx]
                relevant_chunks.append(chunk.page_content)
                sources.add(chunk.metadata.get("source", "Unknown"))
            else:
                print(f"Warning: Invalid chunk index {idx}")
        
        if not relevant_chunks:
            raise HTTPException(status_code=500, detail="No relevant chunks found")
        
        # Create context for LLM
        context = "\n\n".join(relevant_chunks)
        print(f"Context created with {len(relevant_chunks)} chunks")
        
        return context

    except Exception as e:
        print(e)
    return "empty"

@app.post("/message/stream/")
async def send_message_stream(request: MessageRequest):
    """
    Send a message to the current llama.cpp model with streaming response
    """
    history = request.chatHistory
    prompt = request.message.strip()
    ragStatus = request.ragStatus
    session_id= request.session_id
    if not prompt.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")
    
    print(f"Current model: {current_model}")
    print(f'Rag status : {ragStatus}')
    # Generating Chat History
    messages = []
    
    # Add system messages
    messages.append({
        "role": "system",
        "content": "You are a helpful AI assistant 'V'. Provide accurate, concise, and engaging responses. Be conversational and friendly while staying professional. Give direct answers with relevant context. Acknowledge uncertainty rather than guessing. Ask clarifying questions when needed."
    })

    if history and history.strip():
        #messages seperate by \n
        lines = history.strip().split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('V: '):
                messages.append({
                    "role": "assistant",
                    "content": line[3:]  # Remove 'V: '
                })
            elif line.startswith('User: '):
                messages.append({
                    "role": "user",
                    "content": line[6:]  # Remove 'User: '
                })
                
    if ragStatus:
        messages = []
        messages.append({
            "role": "system",
            "content": "You are a helpful AI assistant 'V'. Provide accurate, concise, and engaging responses. Be conversational and friendly while staying professional. Give direct answers with relevant context. Acknowledge uncertainty rather than guessing. Ask clarifying questions when needed."
        })
        context = await gen_context(session_id,prompt)
        messages.append({
                "role": "system", 
                "content": f"DOCUMENT CONTEXT: The user has uploaded a document. Here is relevant information from their document that relates to their question:\n\n{context}\n\nUse this context to inform your response to the user's question."
            })

    # Add current user message
    messages.append({
        "role": "user",
        "content": prompt
    })
    
    print(f"Messages to send: {messages}")
    
    async def generate_stream():
        try:
            # Check if llama.cpp server is running
            async with httpx.AsyncClient(timeout=180.0) as client:
                # Test connection first
                try:
                    await client.get(f"{MODEL_PATH}/health", timeout=5.0)
                except Exception:
                    # Fallback health check if /health doesn't exist
                    try:
                        await client.get(f"{MODEL_PATH}/v1/models", timeout=5.0)
                    except Exception:
                        yield f"data: {json.dumps({'error': 'llama.cpp server is not responding'})}\n\n"
                        return
                
                # Open API standard call with streaming
                print(f"Sending streaming request to llama.cpp server")
                
                async with client.stream(
                    method="POST",
                    url=f"{MODEL_PATH}/v1/chat/completions",
                    json={
                        "model": current_model,
                        "messages": messages,
                        "temperature": 0.3,
                        "top_p": 0.95,
                        "max_tokens": 1000,
                        "stream": True,  # Enable streaming
                        "stop": ["<|eot_id|>", "<|end_of_text|>"]
                    },
                    headers={
                        "Content-Type": "application/json"
                    },
                    timeout=180.0
                ) as response:
                    
                    if response.status_code != 200:
                        print(f"llama.cpp API error: {response.status_code}")
                        yield f"data: {json.dumps({'error': f'llama.cpp API returned status {response.status_code}'})}\n\n"
                        return
                    
                    # Process streaming response
                    async for chunk in response.aiter_lines():
                        if chunk:
                            # Remove 'data: ' prefix if present
                            if chunk.startswith('data: '):
                                chunk = chunk[6:]
                            
                            # Skip empty lines and [DONE] marker
                            if not chunk.strip() or chunk.strip() == '[DONE]':
                                continue
                            
                            try:
                                # Parse the JSON chunk
                                data = json.loads(chunk)
                                
                                # Extract the content from OpenAI-style streaming response
                                if "choices" in data and data["choices"]:
                                    choice = data["choices"][0]
                                    if "delta" in choice and "content" in choice["delta"]:
                                        content = choice["delta"]["content"]
                                        if content:
                                            # Send the token to the frontend
                                            yield f"data: {json.dumps({'token': content})}\n\n"
                                    
                                    # Check if streaming is finished
                                    if choice.get("finish_reason"):
                                        print("finished streaming")
                                        yield f"data: {json.dumps({'done': True})}\n\n"
                                        break
                                        
                            except json.JSONDecodeError as e:
                                print(f"JSON decode error: {e}, chunk: {chunk}")
                                continue
                            except Exception as e:
                                print(f"Error processing chunk: {e}")
                                continue
                
        except HTTPException:
            yield f"data: {json.dumps({'error': 'HTTP exception occurred'})}\n\n"
        except Exception as e:
            print(f"llama.cpp Streaming Error: {str(e)}")
            yield f"data: {json.dumps({'error': f'Error generating response: {str(e)}'})}\n\n"
    
    return StreamingResponse(
        generate_stream(),
        media_type="text/plain",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Content-Type": "text/event-stream",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "Content-Type",
        }
    )

@app.get("/models")
async def get_available_models():
    """Get list of available Ollama models"""
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get("http://localhost:11434/api/tags")
            if response.status_code == 200:
                data = response.json()
                models = [model["name"] for model in data.get("models", [])]
                return JSONResponse({
                    "status": "success",
                    "models": models,
                    "current_model": current_model
                })
            else:
                return JSONResponse({
                    "status": "error",
                    "message": "Failed to fetch models"
                }, status_code=500)
    except Exception as e:
        return JSONResponse({
            "status": "error",
            "message": f"Error fetching models: {str(e)}"
        }, status_code=500)

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            response = await client.get("http://localhost:11434/api/tags")
            if response.status_code != 200:
                return JSONResponse({
                    "status": "unhealthy",
                    "message": "Ollama server not responding"
                }, status_code=503)
            
            data = response.json()
            available_models = [model["name"] for model in data.get("models", [])]
            
            return JSONResponse({
                "status": "healthy",
                "ollama_running": True,
                "current_model": current_model,
                "model_available": current_model in available_models,
                "available_models": available_models,
                "active_rag_sessions": len(rag_sessions)
            })
            
    except Exception as e:
        return JSONResponse({
            "status": "unhealthy",
            "message": f"Health check failed: {str(e)}",
            "ollama_running": False
        }, status_code=503)

# Document processing functions
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    doc = DocxDocument(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

async def get_embeddings(texts: List[str]) -> np.ndarray:
    """Get embeddings from llama.cpp server with Nomic model"""
    # Prepend required prefix to all texts
    prefixed_texts = [f"search_query: {text}" for text in texts]
    
    async with httpx.AsyncClient(timeout=60.0) as client:
        try:
            response = await client.post(
                url=f"{EMBEDDING_PATH}/v1/embeddings",
                json={"input": prefixed_texts},  # Send all texts in one request
                timeout=60.0
            )
            
            if response.status_code == 200:
                data = response.json()
                # Extract embeddings from all items in response data
                return np.array([item["embedding"] for item in data["data"]], dtype=np.float32)
            else:
                raise Exception(f"API error {response.status_code}: {response.text}")
                
        except Exception as e:
            print(f"Embedding error: {e}")
            raise

def create_chunks(text: str, filename: str) -> List[Document]:
    """Split text into chunks"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1024,
        chunk_overlap=20,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    chunks = text_splitter.split_text(text)
    documents = []
    
    for i, chunk in enumerate(chunks):
        doc = Document(
            page_content=chunk,
            metadata={
                "source": filename,
                "chunk_id": i,
                "total_chunks": len(chunks)
            }
        )
        documents.append(doc)
    
    return documents

# RAG Session endpoints
class SessionResponse(BaseModel):
    session_id: str
    status: str
    message: str

@app.post("/rag/create-session", response_model=SessionResponse)
async def create_rag_session():
    """Create a new RAG session"""
    session_id = str(uuid.uuid4())
    rag_sessions[session_id] = RAGSession(session_id)
    
    return SessionResponse(
        session_id=session_id,
        status="success",
        message="RAG session created successfully"
    )

@app.delete("/rag/session/{session_id}")
async def delete_rag_session(session_id: str):
    """Delete a RAG session"""
    if session_id not in rag_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    del rag_sessions[session_id]
    
    return JSONResponse({
        "status": "success",
        "message": "Session deleted successfully"
    })

class RAGUploadResponse(BaseModel):
    status: str
    message: str
    chunks_created: int
    processing_time: float
    filename: str
    session_id: str

@app.post("/rag/upload/{session_id}", response_model=RAGUploadResponse)
async def upload_document_to_session(
    session_id: str,
    file: UploadFile = File(...)
):
    """Upload and process a document for RAG in a specific session"""
    if session_id not in rag_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = rag_sessions[session_id]
    session.update_access_time()
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    # Check if session already has a document
    if session.has_document():
        raise HTTPException(
            status_code=400, 
            detail=f"Session already has a document: {session.document_name}. Remove it first or create a new session."
        )
    
    # Validate file type
    file_extension = os.path.splitext(file.filename)[1].lower()
    supported_extensions = ['.pdf', '.txt', '.docx', '.doc']
    
    if file_extension not in supported_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Supported types: {', '.join(supported_extensions)}"
        )
    
    start_time = time.time()
    temp_file_path = None
    
    try:
        # Save temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text = extract_text_from_pdf(temp_file_path)
        elif file_extension in ['.docx', '.doc']:
            text = extract_text_from_docx(temp_file_path)
        elif file_extension == '.txt':
            text = extract_text_from_txt(temp_file_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text content found in the document")
        
        # Create chunks
        chunks = create_chunks(text, file.filename)
        
        if not chunks:
            raise HTTPException(status_code=400, detail="Failed to create chunks from document")
        
        # Get embeddings
        chunk_texts = [doc.page_content for doc in chunks]
        embeddings = await get_embeddings(chunk_texts)
        
        # Create FAISS index
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatIP(dimension)  # Inner product for similarity
        
        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(embeddings)
        index.add(embeddings)
        
        # Store in session
        session.document_name = file.filename
        session.chunks = chunks
        session.embeddings = embeddings
        session.vector_store = index
        session.chunk_texts = chunk_texts
        
        processing_time = time.time() - start_time
        
        return RAGUploadResponse(
            status="success",
            message=f"Document processed successfully. Created {len(chunks)} chunks.",
            chunks_created=len(chunks),
            processing_time=round(processing_time, 2),
            filename=file.filename,
            session_id=session_id
        )
        
    except HTTPException:
        # Clean up session data on error
        session.clear_document_data()
        raise
    except Exception as e:
        # Clean up session data on error
        session.clear_document_data()
        print(f"Error processing document: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing document: {str(e)}"
        )
    finally:
        # Clean up temporary file
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

@app.delete("/rag/document/{session_id}")
async def remove_document_from_session(session_id: str):
    """Remove document and all processed data from session"""
    if session_id not in rag_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = rag_sessions[session_id]
    session.update_access_time()
    
    if not session.has_document():
        raise HTTPException(status_code=400, detail="No document found in session")
    
    document_name = session.document_name
    session.clear_document_data()
    
    return JSONResponse({
        "status": "success",
        "message": f"Document '{document_name}' and all processed data removed from session"
    })

class RAGQueryRequest(BaseModel):
    query: str
    max_results: Optional[int] = 4

class RAGQueryResponse(BaseModel):
    response: str
    sources: List[str]
    relevant_chunks: List[str]
    status: str
    session_id: str

@app.post("/rag/query/{session_id}")
async def query_rag_session(session_id: str, request: RAGQueryRequest):
    """Query the RAG system for a specific session"""
    if session_id not in rag_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = rag_sessions[session_id]
    session.update_access_time()
    
    if not session.has_document():
        raise HTTPException(
            status_code=400,
            detail="No document found in session. Please upload a document first."
        )
    
    query = request.query.strip()
    if not query:
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    
    try:
        # Get query embedding
        print(f"Getting embedding for query: {query}")
        query_embedding = await get_embeddings([query])
        print(f"Query embedding shape: {query_embedding.shape}")
        
        # Ensure embedding is float32 and normalized
        query_embedding = query_embedding.astype(np.float32)
        faiss.normalize_L2(query_embedding)
        
        # Validate vector store
        if session.vector_store is None:
            raise HTTPException(status_code=500, detail="Vector store not initialized")
        
        # Search similar chunks
        k = min(request.max_results, len(session.chunks))
        print(f"Searching for {k} similar chunks")
        
        # Perform the search
        scores, indices = session.vector_store.search(query_embedding, k)
        print(f"Search completed. Scores: {scores}, Indices: {indices}")
        
        # Get relevant chunks
        relevant_chunks = []
        sources = set()
        
        for i, idx in enumerate(indices[0]):
            if 0 <= idx < len(session.chunks):  # Validate index bounds
                chunk = session.chunks[idx]
                relevant_chunks.append(chunk.page_content)
                sources.add(chunk.metadata.get("source", "Unknown"))
            else:
                print(f"Warning: Invalid chunk index {idx}")
        
        if not relevant_chunks:
            raise HTTPException(status_code=500, detail="No relevant chunks found")
        
        # Create context for LLM
        context = "\n\n".join(relevant_chunks)
        print(f"Context created with {len(relevant_chunks)} chunks")
        
        # Create prompt for LLM
        rag_prompt = f"""You are a helpful AI assistant. Answer the user's question based on the provided context. If the context doesn't contain enough information to answer the question, say so clearly.

CONTEXT:
{context}

QUESTION: {query}

ANSWER: Provide a comprehensive answer based on the context above. If the context is insufficient, clearly state what information is missing."""

    except Exception as e:
        print(f"llama.cpp Streaming Error: {str(e)}")
        
    async def generate_stream():
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                print("Sending prompt to model")
                async with client.stream(
                        method="POST",
                        url=f"{MODEL_PATH}/completion",
                        json={
                            "model": current_model,
                            "prompt": rag_prompt,
                            "temperature": 0.3,
                            "top_p": 0.95,
                            "stream": True,
                            "stop": ["<|eot_id|>", "<|end_of_text|>"]
                        },
                        headers={
                            "Content-Type": "application/json"
                        },
                        timeout=180.0
                    ) as response:
                        
                        if response.status_code != 200:
                            print(f"llama.cpp API error: {response.status_code}")
                            yield f"data: {json.dumps({'error': f'llama.cpp API returned status {response.status_code}'})}\n\n"
                            return
                        
                        # Process streaming response
                        async for chunk in response.aiter_lines():
                            if chunk:
                                # Remove 'data: ' prefix if present
                                if chunk.startswith('data: '):
                                    chunk = chunk[6:]
                                
                                # Skip empty lines and [DONE] marker
                                if not chunk.strip() or chunk.strip() == '[DONE]':
                                    continue
                                
                                try:
                                    # Parse the JSON chunk
                                    data = json.loads(chunk)
                                    
                                    # Extract the content from OpenAI-style streaming response
                                    if "choices" in data and data["choices"]:
                                        choice = data["choices"][0]
                                        if "delta" in choice and "content" in choice["delta"]:
                                            content = choice["delta"]["content"]
                                            if content:
                                                # Send the token to the frontend
                                                yield f"data: {json.dumps({'token': content})}\n\n"
                                        
                                        # Check if streaming is finished
                                        if choice.get("finish_reason"):
                                            yield f"data: {json.dumps({'done': True})}\n\n"
                                            break
                                            
                                except json.JSONDecodeError as e:
                                    print(f"JSON decode error: {e}, chunk: {chunk}")
                                    continue
                                except Exception as e:
                                    print(f"Error processing chunk: {e}")
                                    continue
                    
        except HTTPException:
            yield f"data: {json.dumps({'error': 'HTTP exception occurred'})}\n\n"
        except Exception as e:
            print(f"llama.cpp Streaming Error: {str(e)}")
            yield f"data: {json.dumps({'error': f'Error generating response: {str(e)}'})}\n\n"
    
    return StreamingResponse(
        generate_stream(),
        media_type="text/plain",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Content-Type": "text/event-stream",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "Content-Type",
        }
    )
      
def create_fallback_response(query: str, chunks: List[str]) -> str:
    """Create a simple fallback response when LLM fails"""
    # Simple keyword matching and extraction
    query_words = query.lower().split()
    relevant_sentences = []
    
    for chunk in chunks[:2]:  # Use first 2 chunks
        sentences = chunk.split('.')
        for sentence in sentences:
            if any(word in sentence.lower() for word in query_words):
                relevant_sentences.append(sentence.strip())
                if len(relevant_sentences) >= 3:
                    break
        if len(relevant_sentences) >= 3:
            break
    
    if relevant_sentences:
        return f"Based on the document: {' '.join(relevant_sentences[:3])}"
    else:
        return f"I found relevant information in the document, but couldn't generate a detailed response. The content discusses: {chunks[0][:200]}..."

@app.get("/rag/session/{session_id}/status")
async def get_session_status(session_id: str):
    """Get status of a RAG session"""
    if session_id not in rag_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = rag_sessions[session_id]
    session.update_access_time()
    
    return JSONResponse({
        "session_id": session_id,
        "has_document": session.has_document(),
        "document_name": session.document_name,
        "chunks_count": len(session.chunks),
        "created_at": session.created_at.isoformat(),
        "last_accessed": session.last_accessed.isoformat(),
        "embedding_model": embedding_model,
        "status": "success"
    })

@app.get("/rag/sessions")
async def list_rag_sessions():
    """List all active RAG sessions"""
    await cleanup_expired_sessions()
    
    sessions_info = []
    for session_id, session in rag_sessions.items():
        sessions_info.append({
            "session_id": session_id,
            "has_document": session.has_document(),
            "document_name": session.document_name,
            "chunks_count": len(session.chunks),
            "created_at": session.created_at.isoformat(),
            "last_accessed": session.last_accessed.isoformat()
        })
    
    return JSONResponse({
        "active_sessions": len(rag_sessions),
        "sessions": sessions_info,
        "status": "success"
    })

# File upload (non-RAG)
class FileUploadResponse(BaseModel):
    response: str
    file_info: dict
    status: str

def get_file_type_info(file_path: str, filename: str) -> dict:
    """Get comprehensive file type information"""
    file_info = {}
    
    # Get file extension
    _, extension = os.path.splitext(filename)
    file_info['extension'] = extension.lower() if extension else 'No extension'
    
    # Get MIME type using mimetypes module
    mime_type, _ = mimetypes.guess_type(filename)
    file_info['mime_type'] = mime_type or 'Unknown'
    
    # Get file size
    file_info['size_bytes'] = os.path.getsize(file_path)
    file_info['size_mb'] = round(file_info['size_bytes'] / (1024 * 1024), 2)
    
    return file_info

@app.post("/upload", response_model=FileUploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """Upload and analyze a file, then delete it"""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Create temporary file path
        temp_dir = tempfile.gettempdir()
        safe_filename = "".join(c for c in file.filename if c.isalnum() or c in (' ', '.', '_', '-')).rstrip()
        file_path = os.path.join(temp_dir, f"temp_{safe_filename}")
        
        try:
            # Save the file temporarily
            with open(file_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
            
            # Get file type information
            file_info = get_file_type_info(file_path, file.filename)
            
            # Create response message with file details
            response_message = f"""File Analysis Complete! 📁 File Details\\ Name: {file.filename}\\ Extension: {file_info['extension']}\\ Size:{file_info['size_mb']} MB ({file_info['size_bytes']:,} bytes)"""
            
            return FileUploadResponse(
                response=response_message,
                file_info=file_info,
                status="success"
            )
            
        finally:
            # Always delete the file after processing
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"File deleted: {file_path}")
    
    except Exception as e:
        # Clean up file if it exists and there was an error
        if 'file_path' in locals() and os.path.exists(file_path):
            os.remove(file_path)
        
        print(f"Error processing file: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred while processing the file: {str(e)}"
        )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

"""
for all servers:-
export LD_LIBRARY_PATH=build/ggml/src:build/src/:$LD_LIBRARY_PATH

for embedding server:-
echo 0 | sudo tee /proc/sys/kernel/yama/ptrace_scope

./build/bin/llama-server -m /home/caio/Downloads/Meta-Llama-3.1-8B-Instruct-Q4_K_M.gguf -c 2048 --port 8080 --host 0.0.0.0 --chat-template llama3

./build/bin/llama-server -m /home/caio/Downloads/nomic-embed-text-v1.Q4_K_M.gguf -c 2048 --port 8081 --host 0.0.0.0 --embeddings
"""